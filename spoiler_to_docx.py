from docx import Document
from docx.shared import Cm, Inches
from operator import itemgetter
import json

spoiler_file = "rom/ALttP - VT_er-no-glitches-0.6.1.5_easy-open-dungeons_key-sanity_30yArpDEGV.json"

f=open(spoiler_file, "r")
spoiler_data = json.loads(f.read())

document = Document()

document.add_heading('Spoiler log for seed hash ' + str(spoiler_data['meta']['seed']), 0)
document.add_heading('Entrances',1)

entrance_table = document.add_table(
    rows=1,
    cols=3,
    style='Medium List 1 Accent 1'
)
hdr_cells = entrance_table.rows[0].cells
hdr_cells[0].text = 'Entrance'
hdr_cells[1].text = 'Exit'
hdr_cells[2].text = 'Direction'
for i, entrance in enumerate(spoiler_data['Entrances']):
    row_cells = entrance_table.add_row().cells
    row_cells[0].text = entrance['entrance']
    row_cells[1].text = entrance['exit']
    row_cells[2].text = entrance['direction']

regions = [
	"Light World",
	"Dark World",
	"Caves",
	"Hyrule Castle",
	"Eastern Palace",
	"Desert Palace",
	"Tower of Hera",
	"Agahnims Tower",
	"Palace of Darkness",
	"Thieves Town",
	"Skull Woods",
	"Swamp Palace",
	"Ice Palace",
	"Misery Mire",
	"Turtle Rock",
	"Ganons Tower"
]

document.add_heading('Items',1)
item_table = document.add_table(
    rows=1,
    cols=3,
    style='Medium List 1 Accent 2'
)

itemarr = []
for region in regions:
    for loc in spoiler_data[region]:
        item = spoiler_data[region].get(loc)
        itemarr.append([item,region,loc])

region_table = document.add_table(
    rows=1,
    cols=3,
    style='Medium List 1 Accent 2'
)
hdr_cells = region_table.rows[0].cells
hdr_cells[0].text = 'Item'
hdr_cells[1].text = 'Region'
hdr_cells[2].text = 'Location'

for itemloc in sorted(itemarr, key=itemgetter(0)):
    row_cells = region_table.add_row().cells
    row_cells[0].text = itemloc[0]
    row_cells[1].text = itemloc[1]
    row_cells[2].text = itemloc[2]

document.save('spoiler_log_doc/spoilerlog.docx')